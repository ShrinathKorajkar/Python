import docx, os
from docx.shared import RGBColor

os.chdir(os.path.dirname(__file__))
doc = docx.Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)


def getText(filename):
    doc = docx.Document(filename)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(' ' + para.text)
    return '\n\n'.join(fulltext)


# print(getText('demo.docx'))

# WRITING

doc2 = docx.Document()
doc2.add_paragraph('hello world', 'Title')
para = doc2.add_paragraph('This is 2 paragraph')
doc2.add_paragraph('this is third paragraph')
para.add_run(' this is added to 2 para')
doc2.add_heading('header 0', 0)
doc2.add_heading('header 1', 1)
doc2.add_heading('header 2', 2)
doc2.add_heading('header 3', 3)
doc2.add_paragraph('This is first page')
# doc2.paragraphs[7].runs[0].add_break(docx.text.WD_BREAK.PAGE)
doc2.add_page_break()
doc2.add_paragraph('This is second page')
doc2.add_picture('zophie.png',
                 width=docx.shared.Inches(3),
                 height=docx.shared.Cm(8))
print(doc2.paragraphs[0].style)
# # doc2.paragraphs[0].style = 'Normal'
# # doc2.paragraphs[0].runs[0].style = 'QuoteChar'
doc2.paragraphs[0].runs[0].underline = True
doc2.paragraphs[0].runs[0].font.name = 'Normal'
doc2.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x22, 0x8b, 0x22)

doc2.save('newdoc.docx')